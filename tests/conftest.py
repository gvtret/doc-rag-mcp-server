"""Pytest fixtures shared across the doc-rag test suite.

Two design choices worth knowing:

1. We avoid loading any embedding model in tests. Sprint 2 testing
   targets are the pipeline up to the chunking stage, plus the FAISS
   reconstruct path. For the reconstruct path we build a synthetic
   `IndexFlatIP` straight from numpy vectors, not from a real encoder.

2. We isolate every test that touches the filesystem inside `tmp_path`
   and set `DOC_RAG_ROOT` so the server module reads from there.
"""
from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pytest

_ROOT = Path(__file__).resolve().parents[1]
_SRC = _ROOT / "src"
if _SRC.is_dir():
    sys.path.insert(0, str(_SRC))


# --------------------------------------------------------------------------
# Tiny config helper
# --------------------------------------------------------------------------

_BASE_CONFIG: Dict[str, Any] = {
    "pipeline_version": "test",
    "paths": {
        "sources_incoming": "sources/incoming",
        "sources_archived": "sources/archived",
        "docs_md_dir": "build/docs_md",
        "tables_json_dir": "build/tables_json",
        "chunks_dir": "build/chunks_jsonl",
        "embeddings_dir": "build/embeddings",
        "index_dir": "build/index",
        "manifest_path": "build/manifest.json",
    },
    "parsing": {
        "pdf_backend": "auto",
        "normalize_whitespace": True,
        "min_chars_per_page": 1,
        "edition_year": {
            "from_pdf_metadata": False,
            "filename_regex": None,
            "by_basename": {},
            "by_source_rel_path": {},
            "by_sha256": {},
        },
        "ocr": {"enabled": False},
    },
    "sectioning": {
        "enable_numbered_headings": True,
        "enable_allcaps_headings": True,
        "min_heading_len": 4,
        "max_heading_len": 120,
    },
    "chunking": {
        "target_tokens": 64,
        "overlap_tokens": 8,
        "dedup_similarity_threshold": 0.0,
    },
    "embeddings": {
        "model_name": "test-stub",
        "device": "cpu",
        "batch_size": 4,
        "normalize": True,
    },
    "index": {"backend": "faiss", "metric": "ip", "top_k": 4},
    "server": {"host": "127.0.0.1", "port": 0},
    "sources": {
        "archive_after_ingest": True,
        "incremental_ingest": True,
    },
    "mcp": {"retrieval_mode": "semantic"},
}


def _write_yaml(path: Path, data: Dict[str, Any]) -> None:
    import yaml

    path.write_text(yaml.safe_dump(data, sort_keys=False), encoding="utf-8")


# --------------------------------------------------------------------------
# tmp corpus root
# --------------------------------------------------------------------------


@pytest.fixture
def tmp_corpus_root(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """Create a self-contained corpus root under tmp_path.

    Layout matches what the server/pipeline expect to find under
    DOC_RAG_ROOT. Returns the root path. The config file is at
    config/config.yaml under the root.
    """
    (tmp_path / "sources" / "incoming").mkdir(parents=True)
    (tmp_path / "sources" / "archived").mkdir(parents=True)
    (tmp_path / "build" / "docs_md").mkdir(parents=True)
    (tmp_path / "build" / "tables_json").mkdir(parents=True)
    (tmp_path / "build" / "chunks_jsonl").mkdir(parents=True)
    (tmp_path / "build" / "embeddings").mkdir(parents=True)
    (tmp_path / "build" / "index").mkdir(parents=True)
    (tmp_path / "config").mkdir()
    _write_yaml(tmp_path / "config" / "config.yaml", _BASE_CONFIG)

    monkeypatch.setenv("DOC_RAG_ROOT", str(tmp_path))
    monkeypatch.delenv("DOC_RAG_API_KEY", raising=False)
    return tmp_path


# --------------------------------------------------------------------------
# Synthetic embeddings + FAISS index
# --------------------------------------------------------------------------


@pytest.fixture
def synthetic_embeddings():
    """Return a factory that produces deterministic normalised vectors.

    Usage: `vectors = synthetic_embeddings(n=10, dim=8)`
    The returned `vectors` is a (n, dim) float32 numpy array with unit L2 norm.
    """
    import numpy as np

    def _factory(n: int, dim: int, seed: int = 42):
        rng = np.random.default_rng(seed)
        v = rng.standard_normal(size=(n, dim)).astype("float32")
        norms = np.linalg.norm(v, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        return v / norms

    return _factory


@pytest.fixture
def synthetic_chunks():
    """Return a factory that produces N chunk dicts in the on-disk schema.

    Each chunk has the minimum fields used downstream: chunk_id, doc_id,
    text, source_file.
    """

    def _factory(n: int, *, doc_id: str = "doc-test", source_file: str = "test.pdf"):
        # `chunk_id` uses a single colon so that
        # _rebuild_faiss_after_delete's rsplit(":", 1) correctly extracts
        # the doc_id prefix.
        return [
            {
                "chunk_id": f"{doc_id}:{i:04d}",
                "doc_id": doc_id,
                "text": f"This is chunk {i} of document {doc_id}.",
                "source_file": source_file,
            }
            for i in range(n)
        ]

    return _factory


# --------------------------------------------------------------------------
# Built index helper
# --------------------------------------------------------------------------


@pytest.fixture
def built_corpus(tmp_corpus_root: Path, synthetic_chunks, synthetic_embeddings):
    """Materialise a small but realistic corpus on disk: manifest, chunks,
    embeddings, FAISS index. Useful for tests that need delete/wipe/
    clean-orphans/dedup to operate on a non-empty state.

    Returns a (root, doc_ids) tuple.
    """
    import numpy as np

    try:
        import faiss  # type: ignore
    except Exception:
        pytest.skip("faiss not installed; cannot build index for this test")

    root = tmp_corpus_root

    # Two synthetic documents, three chunks each.
    docs: List[Tuple[str, List[Dict[str, Any]]]] = [
        ("doc-aaa", synthetic_chunks(3, doc_id="doc-aaa", source_file="aaa.pdf")),
        ("doc-bbb", synthetic_chunks(3, doc_id="doc-bbb", source_file="bbb.pdf")),
    ]

    all_chunks: List[Dict[str, Any]] = []
    for _, chunks in docs:
        all_chunks.extend(chunks)

    # Chunks JSONL
    chunks_path = root / "build" / "chunks_jsonl" / "chunks.jsonl"
    with chunks_path.open("w", encoding="utf-8") as f:
        for chunk in all_chunks:
            f.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    # Embeddings + FAISS
    vectors = synthetic_embeddings(n=len(all_chunks), dim=8)
    np.save(root / "build" / "embeddings" / "embeddings.npy", vectors)

    dim = vectors.shape[1]
    index = faiss.IndexFlatIP(dim)
    index.add(vectors)
    faiss.write_index(index, str(root / "build" / "index" / "faiss.index"))

    # index_meta.json layout matches what _rebuild_faiss_after_delete reads.
    meta = {
        "chunk_ids": [c["chunk_id"] for c in all_chunks],
        "dim": int(dim),
        "metric": "ip",
    }
    (root / "build" / "index" / "index_meta.json").write_text(
        json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    # Source files
    archived = root / "sources" / "archived"
    archived.mkdir(parents=True, exist_ok=True)
    for doc_id, chunks in docs:
        src = archived / chunks[0]["source_file"]
        src.write_bytes(b"%PDF-1.4 stub\n")

    # Markdown copies under build/docs_md/
    md_dir = root / "build" / "docs_md"
    md_dir.mkdir(parents=True, exist_ok=True)
    for doc_id, _ in docs:
        (md_dir / f"{doc_id}.md").write_text(
            f"# {doc_id}\n\nplaceholder\n", encoding="utf-8"
        )

    # Manifest
    manifest = {
        "documents": [
            {
                "doc_id": doc_id,
                "source_file": f"sources/archived/{chunks[0]['source_file']}",
                "md_path": f"build/docs_md/{doc_id}.md",
                "sha256": f"deadbeef{i:056x}",
                "chunks": len(chunks),
            }
            for i, (doc_id, chunks) in enumerate(docs)
        ]
    }
    (root / "build" / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    return root, [doc_id for doc_id, _ in docs]


# --------------------------------------------------------------------------
# Sample document factories
# --------------------------------------------------------------------------


@pytest.fixture
def make_md(tmp_path: Path):
    """Drop a .md file with given text in tmp_path and return its path."""

    def _factory(name: str, text: str) -> Path:
        p = tmp_path / name
        p.write_text(text, encoding="utf-8")
        return p

    return _factory


@pytest.fixture
def make_txt(tmp_path: Path):
    def _factory(name: str, text: str) -> Path:
        p = tmp_path / name
        p.write_text(text, encoding="utf-8")
        return p

    return _factory


@pytest.fixture
def make_docx(tmp_path: Path):
    """Build a real .docx with python-docx; skip the test if it is unavailable."""

    def _factory(name: str, paragraphs: List[str], table_rows: Optional[List[List[str]]] = None) -> Path:
        try:
            from docx import Document
        except Exception:
            pytest.skip("python-docx not installed")
        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        if table_rows:
            rows = len(table_rows)
            cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=rows, cols=cols)
            for i, row in enumerate(table_rows):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text
        p = tmp_path / name
        doc.save(str(p))
        return p

    return _factory


@pytest.fixture
def antiword_available() -> bool:
    return shutil.which("antiword") is not None or shutil.which("catdoc") is not None


@pytest.fixture
def tesseract_available() -> bool:
    return shutil.which("tesseract") is not None
