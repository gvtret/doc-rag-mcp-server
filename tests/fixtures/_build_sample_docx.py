"""Build tests/fixtures/sample.docx — a rich-but-plausible fixture document.

Run from the repo root:

    .venv/bin/python tests/fixtures/_build_sample_docx.py

The generated .docx contains:
  - Heading 1, 2, 3 cascade
  - Plain paragraphs
  - Bold/italic inline runs
  - A bulleted list
  - A numbered list
  - A real 4×3 table with numeric content
  - A small embedded PNG (a schematic) generated on the fly via Pillow

Markers used by tests (kept short and plausible-looking):
  - "АВ-12"           — equipment series mentioned several times
  - "РК-22.04-2024"   — registration code unique enough to assert on

This file generates the .docx; the .docx itself is committed to the
repository, so contributors do not need Pillow or python-docx just to
run the test suite.
"""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


_OUT_DOCX = Path(__file__).resolve().parent / "sample.docx"


def _build_schematic_png() -> bytes:
    """Produce a small monochrome schematic (~12 kB)."""
    w, h = 640, 320
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)

    # Frame
    d.rectangle([10, 10, w - 10, h - 10], outline="black", width=2)

    # Three component boxes representing the measurement bench:
    #   [Source]──→[DUT]──→[Meter]
    box_w, box_h = 140, 80
    y = (h - box_h) // 2
    x_positions = [40, (w - box_w) // 2, w - 40 - box_w]
    labels = ["Источник", "АВ-12", "Измеритель"]

    for x, label in zip(x_positions, labels):
        d.rectangle([x, y, x + box_w, y + box_h], outline="black", width=2)
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 18)
        except OSError:
            font = ImageFont.load_default()
        bbox = d.textbbox((0, 0), label, font=font)
        tx = x + (box_w - (bbox[2] - bbox[0])) // 2
        ty = y + (box_h - (bbox[3] - bbox[1])) // 2
        d.text((tx, ty), label, fill="black", font=font)

    # Connect them with arrows
    for i in range(len(x_positions) - 1):
        x_start = x_positions[i] + box_w
        x_end = x_positions[i + 1]
        y_mid = y + box_h // 2
        d.line([x_start, y_mid, x_end, y_mid], fill="black", width=2)
        # Arrowhead
        d.polygon(
            [(x_end, y_mid), (x_end - 10, y_mid - 6), (x_end - 10, y_mid + 6)],
            fill="black",
        )

    # Caption under the diagram
    caption = "Рис. 1. Схема входного контроля АВ-12"
    try:
        cf = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    except OSError:
        cf = ImageFont.load_default()
    cbox = d.textbbox((0, 0), caption, font=cf)
    cx = (w - (cbox[2] - cbox[0])) // 2
    d.text((cx, h - 40), caption, fill="black", font=cf)

    buf = io.BytesIO()
    img.save(buf, format="PNG", optimize=True)
    return buf.getvalue()


def _add_h(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def build() -> Path:
    doc = Document()

    # ----- Heading and body -----
    _add_h(doc, "Регламент входного контроля автоматических выключателей серии АВ-12", level=0)

    p = doc.add_paragraph()
    p.add_run("Настоящий регламент устанавливает порядок проведения входного контроля "
              "автоматических выключателей серии ").bold = False
    p.add_run("АВ-12").bold = True
    p.add_run(" при поступлении партий на склад комплектующих изделий энергетического "
              "хозяйства предприятия. Регистрационный шифр регламента: ").bold = False
    p.add_run("РК-22.04-2024").bold = True
    p.add_run(".")

    _add_h(doc, "1. Область применения", level=1)
    doc.add_paragraph(
        "Регламент распространяется на партии серии АВ-12 номиналом от 6 до 63 А, "
        "поступающие от поставщиков, прошедших процедуру квалификации в соответствии "
        "с действующим стандартом качества."
    )

    _add_h(doc, "2. Нормативные ссылки", level=1)
    bullets = [
        "ГОСТ 32395-2013. Выключатели автоматические.",
        "ГОСТ Р 50345-2010. Аппаратура малогабаритная электрическая.",
        "Инструкция по комплектации складских остатков, форма ИК-СО-22.04.",
    ]
    for line in bullets:
        doc.add_paragraph(line, style="List Bullet")

    _add_h(doc, "3. Общие положения", level=1)
    p = doc.add_paragraph()
    p.add_run("Контроль выполняется в нормальных климатических условиях по ").bold = False
    p.add_run("ГОСТ 15150").italic = True
    p.add_run(": температура воздуха от +15 до +25 °C, "
              "относительная влажность не более 80 %, атмосферное давление в диапазоне "
              "84-106 кПа.").bold = False

    _add_h(doc, "4. Перечень проверок", level=1)
    _add_h(doc, "4.1. Внешний осмотр", level=2)
    steps = [
        "Проверить целостность корпуса и отсутствие механических повреждений.",
        "Осмотреть клеммы и контактные соединения на предмет окисления.",
        "Убедиться в чёткости и читаемости заводской маркировки.",
        "Сверить комплектность партии с упаковочной ведомостью.",
    ]
    for line in steps:
        doc.add_paragraph(line, style="List Number")

    _add_h(doc, "4.2. Соответствие маркировки", level=2)
    doc.add_paragraph("Таблица 1. Соответствие маркировки требованиям документации.")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Light Grid Accent 1"
    headers = ("Параметр", "Требование", "Допуск")
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    rows = [
        ("Номинальный ток", "16 А", "± 0 %"),
        ("Номинальное напряжение", "230 В", "± 5 %"),
        ("Категория применения", "АС-22А", "строго"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.cell(i, j).text = val

    _add_h(doc, "4.3. Измерение электрических параметров", level=2)
    doc.add_paragraph(
        "Измерения проводятся по схеме, приведённой на рисунке 1. Источник питания "
        "подключается к входным клеммам испытуемого выключателя; на выходных клеммах "
        "снимаются показания измерительного прибора."
    )

    schematic = _build_schematic_png()
    img_buf = io.BytesIO(schematic)
    doc.add_picture(img_buf, width=Cm(14))

    _add_h(doc, "5. Оформление результатов", level=1)
    p = doc.add_paragraph()
    p.add_run("По итогам входного контроля составляется акт по форме ").bold = False
    p.add_run("РК-22.04-2024").bold = True
    p.add_run(". Акт регистрируется в журнале входного контроля и хранится "
              "не менее пяти лет с даты подписания.").bold = False

    _add_h(doc, "6. Ответственность", level=1)
    doc.add_paragraph(
        "Ответственным за проведение входного контроля назначается старший мастер "
        "участка комплектации, утверждаемый приказом по подразделению."
    )

    doc.save(str(_OUT_DOCX))
    return _OUT_DOCX


if __name__ == "__main__":
    out = build()
    size_kb = out.stat().st_size / 1024
    print(f"wrote {out} ({size_kb:.1f} kB)")
